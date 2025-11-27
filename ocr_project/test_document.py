#!/usr/bin/env python3
"""
Test script for document processing functionality
"""
import os
from docx import Document
from ocr.document_processor import DocumentProcessor

def create_test_word_document():
    """Create a test Word document with text and table"""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document for Processing', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is a test document created to verify the document processing functionality.')
    doc.add_paragraph('It contains various types of content including text, tables, and metadata.')
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Fill the table
    cells = table.rows[0].cells
    cells[0].text = 'Name'
    cells[1].text = 'Age'
    cells[2].text = 'City'
    
    cells = table.rows[1].cells
    cells[0].text = 'John Doe'
    cells[1].text = '30'
    cells[2].text = 'New York'
    
    cells = table.rows[2].cells
    cells[0].text = 'Jane Smith'
    cells[1].text = '25'
    cells[2].text = 'Los Angeles'
    
    # Save the document
    filepath = os.path.join('temp_uploads', 'test_document.docx')
    os.makedirs('temp_uploads', exist_ok=True)
    doc.save(filepath)
    return filepath

def test_document_processing():
    """Test the document processing functionality"""
    
    # Create test document
    print("Creating test Word document...")
    doc_path = create_test_word_document()
    print(f"✅ Created test document: {doc_path}")
    
    # Initialize processor
    print("\nInitializing document processor...")
    processor = DocumentProcessor()
    print("✅ Document processor initialized")
    
    # Process the document
    print(f"\nProcessing document: {doc_path}")
    try:
        result = processor.process_document(doc_path)
        
        if not result.get('success', False):
            print(f"❌ Document processing failed: {result.get('error', 'Unknown error')}")
            return False
        
        print("📄 Document Processing Results:")
        print(f"   Format: {result['file_format']}")
        print(f"   Text length: {len(result['text'])} characters")
        print(f"   Word count: {result.get('word_count', 'N/A')}")
        print(f"   Number of tables: {result.get('table_count', 0)}")
        
        # Display first 200 characters of text
        print(f"\n📝 Text preview:")
        print(f"   {result['text'][:200]}...")
        
        # Display table data
        if result.get('tables'):
            print(f"\n📊 Table data:")
            for table_info in result['tables']:
                df = table_info['dataframe']
                print(f"   Table {table_info['table_index']+1}: {df.shape[0]} rows x {df.shape[1]} columns")
                print(f"   Columns: {list(df.columns)}")
                if not df.empty:
                    print(f"   First row: {df.iloc[0].to_dict()}")
        
        # Display metadata
        if result.get('metadata'):
            print(f"\n📋 Metadata:")
            for key, value in result['metadata'].items():
                if value:  # Only show non-empty metadata
                    print(f"   {key}: {value}")
                
        print("\n✅ Document processing test completed successfully!")
        
    except Exception as e:
        print(f"❌ Error processing document: {e}")
        return False
        
    return True

if __name__ == "__main__":
    print("🧪 Starting Document Processing Test")
    print("=" * 50)
    
    success = test_document_processing()
    
    print("=" * 50)
    if success:
        print("🎉 All tests passed! Document processing system is working correctly.")
    else:
        print("❌ Tests failed. Check the error messages above.")