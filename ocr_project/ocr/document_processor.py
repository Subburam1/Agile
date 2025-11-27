"""
Advanced Document Processing Module
Comprehensive support for Word, PDF, and XML document processing
Includes text extraction, table processing, and metadata extraction
"""

import os
import logging
import pandas as pd
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
import xml.etree.ElementTree as ET

# Document processing imports
try:
    import docx
    from docx import Document
    from docx.table import Table as DocxTable
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import lxml.etree as lxml_etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Advanced document processor supporting Word, PDF, and XML formats.
    Provides text extraction, table processing, and metadata extraction.
    """
    
    def __init__(self):
        """Initialize document processor with available libraries."""
        self.supported_formats = []
        
        if DOCX_AVAILABLE:
            self.supported_formats.append('.docx')
        if PYMUPDF_AVAILABLE or PDFPLUMBER_AVAILABLE:
            self.supported_formats.append('.pdf')
        if LXML_AVAILABLE:
            self.supported_formats.extend(['.xml', '.xhtml', '.html'])
        
        logger.info(f"✅ DocumentProcessor initialized with support for: {', '.join(self.supported_formats)}")
        
        # Log available libraries
        available_libs = []
        if DOCX_AVAILABLE:
            available_libs.append("python-docx")
        if PYMUPDF_AVAILABLE:
            available_libs.append("PyMuPDF")
        if PDFPLUMBER_AVAILABLE:
            available_libs.append("pdfplumber")
        if PYPDF2_AVAILABLE:
            available_libs.append("PyPDF2")
        if LXML_AVAILABLE:
            available_libs.append("lxml")
        
        logger.info(f"📚 Available libraries: {', '.join(available_libs)}")
    
    def process_document(self, file_path: str, extract_tables: bool = True, 
                        extract_metadata: bool = True) -> Dict[str, Any]:
        """
        Process a document file and extract text, tables, and metadata.
        
        Args:
            file_path: Path to the document file
            extract_tables: Whether to extract tables as DataFrames
            extract_metadata: Whether to extract document metadata
        
        Returns:
            Dictionary containing extracted content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported: {self.supported_formats}")
        
        try:
            if file_extension == '.docx':
                return self._process_word_document(file_path, extract_tables, extract_metadata)
            elif file_extension == '.pdf':
                return self._process_pdf_document(file_path, extract_tables, extract_metadata)
            elif file_extension in ['.xml', '.xhtml', '.html']:
                return self._process_xml_document(file_path, extract_tables, extract_metadata)
            else:
                raise ValueError(f"Handler not implemented for: {file_extension}")
                
        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_path': str(file_path),
                'file_format': file_extension
            }
    
    def _process_word_document(self, file_path: Path, extract_tables: bool, 
                             extract_metadata: bool) -> Dict[str, Any]:
        """Process Word document using python-docx."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        doc = Document(file_path)
        
        # Extract text content
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        full_text = '\n\n'.join(paragraphs)
        
        result = {
            'success': True,
            'file_path': str(file_path),
            'file_format': '.docx',
            'text': full_text,
            'paragraphs': paragraphs,
            'paragraph_count': len(paragraphs),
            'character_count': len(full_text),
            'word_count': len(full_text.split())
        }
        
        # Extract tables if requested
        if extract_tables:
            tables_data = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    # Convert to DataFrame
                    df = pd.DataFrame(table_data[1:], columns=table_data[0]) if len(table_data) > 1 else pd.DataFrame(table_data)
                    tables_data.append({
                        'table_index': i,
                        'dataframe': df,
                        'raw_data': table_data,
                        'shape': df.shape
                    })
            
            result['tables'] = tables_data
            result['table_count'] = len(tables_data)
        
        # Extract metadata if requested
        if extract_metadata:
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'category': core_props.category or '',
                'comments': core_props.comments or ''
            }
            result['metadata'] = metadata
        
        return result
    
    def _process_pdf_document(self, file_path: Path, extract_tables: bool, 
                            extract_metadata: bool) -> Dict[str, Any]:
        """Process PDF document using PyMuPDF and pdfplumber."""
        result = {
            'success': True,
            'file_path': str(file_path),
            'file_format': '.pdf',
            'pages': []
        }
        
        # Use PyMuPDF for text extraction and metadata
        if PYMUPDF_AVAILABLE:
            doc = fitz.open(file_path)
            
            # Extract text from each page
            all_text = []
            page_data = []
            
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                all_text.append(page_text)
                
                page_info = {
                    'page_number': page_num + 1,
                    'text': page_text,
                    'character_count': len(page_text),
                    'word_count': len(page_text.split())
                }
                page_data.append(page_info)
            
            result['pages'] = page_data
            result['page_count'] = doc.page_count
            result['text'] = '\n\n'.join(all_text)
            result['character_count'] = len(result['text'])
            result['word_count'] = len(result['text'].split())
            
            # Extract metadata
            if extract_metadata:
                metadata = doc.metadata
                result['metadata'] = {
                    'title': metadata.get('title', ''),
                    'author': metadata.get('author', ''),
                    'subject': metadata.get('subject', ''),
                    'keywords': metadata.get('keywords', ''),
                    'creator': metadata.get('creator', ''),
                    'producer': metadata.get('producer', ''),
                    'created': metadata.get('creationDate', ''),
                    'modified': metadata.get('modDate', '')
                }
            
            doc.close()
        
        # Use pdfplumber for table extraction
        if extract_tables and PDFPLUMBER_AVAILABLE:
            tables_data = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    
                    for table_index, table in enumerate(page_tables):
                        if table and len(table) > 0:
                            # Convert to DataFrame
                            df = pd.DataFrame(table[1:], columns=table[0]) if len(table) > 1 else pd.DataFrame(table)
                            
                            tables_data.append({
                                'page_number': page_num + 1,
                                'table_index': table_index,
                                'dataframe': df,
                                'raw_data': table,
                                'shape': df.shape
                            })
            
            result['tables'] = tables_data
            result['table_count'] = len(tables_data)
        
        # Fallback to PyPDF2 for basic extraction if PyMuPDF not available
        elif PYPDF2_AVAILABLE:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                all_text = []
                page_data = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    all_text.append(page_text)
                    
                    page_info = {
                        'page_number': page_num + 1,
                        'text': page_text,
                        'character_count': len(page_text),
                        'word_count': len(page_text.split())
                    }
                    page_data.append(page_info)
                
                result['pages'] = page_data
                result['page_count'] = len(pdf_reader.pages)
                result['text'] = '\n\n'.join(all_text)
                result['character_count'] = len(result['text'])
                result['word_count'] = len(result['text'].split())
                
                # Extract metadata
                if extract_metadata and pdf_reader.metadata:
                    metadata = pdf_reader.metadata
                    result['metadata'] = {
                        'title': metadata.get('/Title', ''),
                        'author': metadata.get('/Author', ''),
                        'subject': metadata.get('/Subject', ''),
                        'keywords': metadata.get('/Keywords', ''),
                        'creator': metadata.get('/Creator', ''),
                        'producer': metadata.get('/Producer', ''),
                        'created': str(metadata.get('/CreationDate', '')),
                        'modified': str(metadata.get('/ModDate', ''))
                    }
        
        return result
    
    def _process_xml_document(self, file_path: Path, extract_tables: bool, 
                            extract_metadata: bool) -> Dict[str, Any]:
        """Process XML document using lxml or xml.etree.ElementTree."""
        
        # Try lxml first for better performance
        if LXML_AVAILABLE:
            try:
                tree = lxml_etree.parse(str(file_path))
                root = tree.getroot()
                parser_used = 'lxml'
            except Exception as e:
                logger.warning(f"lxml parsing failed, falling back to ElementTree: {e}")
                tree = ET.parse(file_path)
                root = tree.getroot()
                parser_used = 'xml.etree.ElementTree'
        else:
            tree = ET.parse(file_path)
            root = tree.getroot()
            parser_used = 'xml.etree.ElementTree'
        
        # Extract all text content
        all_text = []
        elements_data = []
        
        def extract_element_data(element, path=""):
            """Recursively extract data from XML elements."""
            current_path = f"{path}/{element.tag}" if path else element.tag
            
            element_info = {
                'tag': element.tag,
                'path': current_path,
                'text': element.text.strip() if element.text else '',
                'attributes': dict(element.attrib),
                'children_count': len(list(element))
            }
            
            if element_info['text']:
                all_text.append(element_info['text'])
            
            elements_data.append(element_info)
            
            # Process children
            for child in element:
                extract_element_data(child, current_path)
        
        extract_element_data(root)
        
        result = {
            'success': True,
            'file_path': str(file_path),
            'file_format': file_path.suffix.lower(),
            'parser_used': parser_used,
            'root_tag': root.tag,
            'text': '\n'.join(all_text),
            'elements': elements_data,
            'element_count': len(elements_data),
            'character_count': len('\n'.join(all_text)),
            'word_count': len(' '.join(all_text).split())
        }
        
        # Extract table-like structures if requested
        if extract_tables:
            tables_data = []
            
            # Look for common table-like structures
            table_tags = ['table', 'grid', 'list', 'items', 'rows']
            
            for table_tag in table_tags:
                # Find all elements with table-like names
                if LXML_AVAILABLE:
                    table_elements = tree.xpath(f".//{table_tag}")
                else:
                    table_elements = root.findall(f".//{table_tag}")
                
                for i, table_element in enumerate(table_elements):
                    # Extract table data
                    rows = []
                    
                    # Look for row-like children
                    for child in table_element:
                        if child.tag in ['row', 'item', 'tr', 'record']:
                            row_data = []
                            
                            # Extract cell data
                            for cell in child:
                                cell_text = cell.text.strip() if cell.text else ''
                                row_data.append(cell_text)
                            
                            if row_data:
                                rows.append(row_data)
                    
                    if rows:
                        # Convert to DataFrame
                        df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
                        
                        tables_data.append({
                            'table_tag': table_tag,
                            'table_index': i,
                            'dataframe': df,
                            'raw_data': rows,
                            'shape': df.shape,
                            'xpath': f".//{table_tag}[{i+1}]"
                        })
            
            result['tables'] = tables_data
            result['table_count'] = len(tables_data)
        
        # Extract metadata from root attributes
        if extract_metadata:
            metadata = {
                'root_tag': root.tag,
                'root_attributes': dict(root.attrib),
                'namespace': root.tag.split('}')[0][1:] if '}' in root.tag else '',
                'encoding': tree.docinfo.encoding if hasattr(tree, 'docinfo') else 'unknown'
            }
            
            # Look for common metadata elements
            metadata_tags = ['title', 'author', 'description', 'created', 'modified', 'version']
            for tag in metadata_tags:
                if LXML_AVAILABLE:
                    elements = tree.xpath(f".//{tag}")
                else:
                    elements = root.findall(f".//{tag}")
                
                if elements:
                    metadata[tag] = elements[0].text.strip() if elements[0].text else ''
            
            result['metadata'] = metadata
        
        return result
    
    def extract_text_only(self, file_path: str) -> str:
        """Extract only text content from a document."""
        result = self.process_document(file_path, extract_tables=False, extract_metadata=False)
        return result.get('text', '') if result.get('success') else ''
    
    def extract_tables_only(self, file_path: str) -> List[pd.DataFrame]:
        """Extract only tables from a document as DataFrames."""
        result = self.process_document(file_path, extract_tables=True, extract_metadata=False)
        
        if result.get('success') and 'tables' in result:
            return [table['dataframe'] for table in result['tables']]
        
        return []
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic information about a document."""
        file_path = Path(file_path)
        
        info = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_extension': file_path.suffix.lower(),
            'file_size_bytes': file_path.stat().st_size if file_path.exists() else 0,
            'supported': file_path.suffix.lower() in self.supported_formats
        }
        
        if info['file_size_bytes'] > 0:
            info['file_size_mb'] = round(info['file_size_bytes'] / (1024 * 1024), 2)
        
        return info


def create_document_processor() -> DocumentProcessor:
    """Factory function to create a DocumentProcessor instance."""
    return DocumentProcessor()


def process_document_file(file_path: str, **kwargs) -> Dict[str, Any]:
    """
    Convenience function to process a document file.
    
    Args:
        file_path: Path to the document file
        **kwargs: Additional arguments passed to process_document()
    
    Returns:
        Dictionary containing extracted content and metadata
    """
    processor = DocumentProcessor()
    return processor.process_document(file_path, **kwargs)


def get_supported_formats() -> List[str]:
    """Get list of supported document formats."""
    processor = DocumentProcessor()
    return processor.supported_formats


# Export main functions and classes
__all__ = [
    'DocumentProcessor',
    'create_document_processor', 
    'process_document_file',
    'get_supported_formats'
]